from docx import Document
import re
import unicodedata
from config import image_data
from tkinter import messagebox
import tkinter as tk
import os
import win32com.client as win32

def normalize_bookmark_name(name):
    name = re.sub(r"[\s']", "_", name)
    name = unicodedata.normalize('NFKD', name).encode('ascii', 'ignore').decode('utf-8')
    return name

def trouver_orphelins_initiaux(doc_path):
    """Retourne un ensemble de signets orphelins en cascade en vérifiant les dépendances initiales dans `image_data`."""
    document = demander_document_word(doc_path)
    if not document:
        return set()

    rouge_cascade = set()
    non_placables = set()
    candidats_orphelins = []

    for item in image_data:
        target_name = item['target']

        if not is_word_target_found(target_name, document):
            candidats_orphelins.append(normalize_bookmark_name(target_name))

    for target in candidats_orphelins:
        orphelin_detecte = False  
        
        if target in non_placables:
            for item in image_data:
                if normalize_bookmark_name(item['target']) == target:
                    non_placables.add(normalize_bookmark_name(item['bookmark_name']))
                    orphelin_detecte = True
                    break

        else:
            trouve = False
            for item in image_data:
                bookmark_name = normalize_bookmark_name(item['bookmark_name'])
                
                if target == bookmark_name:
                    trouve = True
                    break
            
            if not trouve:
                for item in image_data:
                    if normalize_bookmark_name(item['target']) == target:
                        non_placables.add(normalize_bookmark_name(item['bookmark_name']))
                        orphelin_detecte = True

        if orphelin_detecte:
            rouge_cascade.update(non_placables)

    return rouge_cascade

def demander_document_word(doc_path):
    """Retourne le Document Word chargé ou None si non trouvé."""
    try:
        if doc_path:
            return Document(doc_path)
    except Exception as e:
        messagebox.showerror("Erreur", f"Erreur lors du chargement du document Word : {e}")
    return None

def is_word_target_found(target, document):
    """Vérifie si le texte `target` est présent dans le document Word, en adaptant le format si nécessaire."""
    if not document:
        return False

    # Appliquer la conversion seulement si `target` est entièrement en majuscules
    if target.isupper():
        formatted_target = target.lower().capitalize()
    else:
        formatted_target = target  # Utiliser `target` tel quel si ce n'est pas tout en majuscules

    # Parcourir chaque paragraphe pour vérifier la présence de `formatted_target`
    for paragraph in document.paragraphs:
        if formatted_target in paragraph.text:
            return True

    return False

def display_paragraph_selection(doc_path, images_folder, image_name, new_card_window, maps_window):
    """Affiche une fenêtre pour sélectionner l'emplacement où insérer le repère de la nouvelle carte."""
    if not image_name:
        messagebox.showwarning("Informations manquantes", "Veuillez entrer le nom de l'image.")
        return

    # Charger le document Word avec python-docx
    try:
        doc = Document(doc_path)
    except Exception as e:
        messagebox.showerror("Erreur", f"Impossible de charger le document Word : {e}")
        return

    # Vérifier que l'image n'existe pas déjà dans le document
    repere = f"<<{image_name}>>"
    for paragraph in doc.paragraphs:
        if repere in paragraph.text:
            messagebox.showerror("Erreur", f"Un repère '{repere}' existe déjà dans le document.")
            return

    # Vérifier que l'image existe dans le dossier des images
    chemin_image = os.path.join(images_folder, f"{image_name}.jpg")
    if not os.path.exists(chemin_image):
        messagebox.showerror("Erreur", f"L'image '{image_name}.jpg' est introuvable dans le dossier des images.")
        return

    # Fenêtre pour sélectionner l'emplacement du repère
    selection_window = tk.Toplevel(new_card_window)
    selection_window.title("Sélectionner l'emplacement pour la carte")
    selection_window.geometry("500x600")
    selection_window.config(bg="#f5f5f5")

    tk.Label(selection_window, text="Choisissez le paragraphe ou le signet pour placer le repère :", bg="#f5f5f5", font=("Helvetica", 12)).pack(pady=10)

    # Ajout du Canvas et de la barre de défilement pour la liste de paragraphes et signets
    canvas = tk.Canvas(selection_window, bg="#f5f5f5")
    scrollbar = tk.Scrollbar(selection_window, orient="vertical", command=canvas.yview)
    scrollable_frame = tk.Frame(canvas, bg="#f5f5f5")

    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )
    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)

    # Pack Canvas et Scrollbar
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Ajouter les paragraphes avec boutons de sélection
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():  # Afficher seulement les paragraphes non vides
            frame = tk.Frame(scrollable_frame, bg="#f5f5f5")
            frame.pack(fill="x", padx=10, pady=2)
            tk.Label(frame, text=paragraph.text[:100], bg="#f5f5f5", wraplength=400, anchor="w").pack(side="left")
            tk.Button(frame, text="Sélectionner", command=lambda text=paragraph.text: ajouter_repere(
                doc_path, repere, chemin_image, bookmark_name=text.strip(), is_paragraph=True,
                new_card_window=new_card_window, maps_window=maps_window, selection_window=selection_window)).pack(side="right")


    # Ajouter les signets avec boutons de sélection en utilisant win32com
    word_app = win32.Dispatch("Word.Application")
    word_app.Visible = False
    document = word_app.Documents.Open(os.path.normpath(doc_path))
    
    for bookmark in document.Bookmarks:
        frame = tk.Frame(scrollable_frame, bg="#f5f5f5")
        frame.pack(fill="x", padx=10, pady=2)
        tk.Label(frame, text=f"Signet : {bookmark.Name}", bg="#f5f5f5", font=("Helvetica", 10, "italic")).pack(side="left")
        tk.Button(frame, text="Sélectionner", command=lambda bookmark=bookmark.Name: ajouter_repere(
            doc_path, repere, chemin_image, bookmark_name=bookmark if isinstance(bookmark, str) else bookmark.Name, is_paragraph=False,
            new_card_window=new_card_window, maps_window=maps_window, selection_window=selection_window)).pack(side="right")
    
    document.Close(SaveChanges=False)
    word_app.Quit()

def ajouter_repere(doc_path, repere, chemin_image, bookmark_name=None, is_paragraph=True, new_card_window=None, maps_window=None, selection_window=None):
    """Ajoute un repère en appelant la fonction d'insertion de repère dans le document Word."""
    # Vérifier que l'image existe avant d'ajouter le repère
    if not os.path.exists(chemin_image):
        messagebox.showwarning("Image manquante", f"L'image '{chemin_image}' est introuvable.")
        return

    # Appel de la fonction qui insère le repère dans le document Word
    ajouter_repere_dans_doc(doc_path, repere, chemin_image, bookmark_name=bookmark_name, is_paragraph=is_paragraph)

    # Mise à jour de l'interface graphique si nécessaire
    if new_card_window:
        new_card_window.destroy()
    if selection_window:
        selection_window.destroy()
    if maps_window:
        maps_window.update()  # Mettre à jour la fenêtre principale pour refléter l'ajout du repère

def ajouter_repere_dans_doc(doc_path, repere, chemin_image, bookmark_name=None, is_paragraph=True):
    """Ajoute un repère sous forme de texte dans le document Word."""
    word_app = win32.Dispatch("Word.Application")
    word_app.Visible = False
    document = word_app.Documents.Open(os.path.normpath(doc_path))
    
    try:
        
        # Insérer le repère après le paragraphe ou le signet sélectionné
        if is_paragraph:
            # Rechercher le paragraphe correspondant au texte donné
            range_emplacement = None
            for paragraph in document.Paragraphs:
                if paragraph.Range.Text.strip() == bookmark_name.strip():
                    range_emplacement = paragraph.Range
                    break
            if not range_emplacement:
                messagebox.showwarning("Paragraphe introuvable", f"Le paragraphe spécifié '{bookmark_name}' est introuvable.")
                return
        else:
            # Ajouter le repère après le signet
            if bookmark_name in document.Bookmarks:
                range_emplacement = document.Bookmarks(bookmark_name).Range
            else:
                messagebox.showwarning("Signet introuvable", f"Le signet '{bookmark_name}' est introuvable.")
                return
        
        range_emplacement.InsertAfter("\r\n")
        range_emplacement.InsertAfter(repere)
        new_range = document.Range(range_emplacement.End, range_emplacement.End)
        new_range.InsertParagraphAfter()
        new_range.InsertParagraphAfter()  # Double saut de paragraphe pour s'assurer de l'espace

        # Appliquer le style de texte
        repere_range = range_emplacement.Paragraphs.Last.Range
        repere_range.Font.Size = 1
        repere_range.Font.Color = 16777215  # Blanc
        repere_range.ParagraphFormat.Alignment = 0

        # Sauvegarder le document
        document.Save()
    finally:
        document.Close(SaveChanges=True)
        word_app.Quit()

    messagebox.showinfo("Carte ajoutée", f"Le repère '{repere}' a été ajouté avec succès.")

def ajouter_nouvelle_carte(doc_path, images_folder, maps_window):
    """Ouvre une fenêtre pour ajouter une nouvelle carte avec un texte repère."""
    # Demander seulement le nom de l'image
    new_card_window = tk.Toplevel(maps_window)
    new_card_window.title("Ajouter une nouvelle carte")
    new_card_window.geometry("400x300")
    new_card_window.config(bg="#f5f5f5")

    tk.Label(new_card_window, text="Nom de l'image (sans extension):", bg="#f5f5f5").pack(pady=10)
    image_name_entry = tk.Entry(new_card_window)
    image_name_entry.pack(pady=5)

    # Bouton pour sélectionner l'emplacement
    emplacement_button = tk.Button(new_card_window, text="Sélectionner l'emplacement",
                                   command=lambda: display_paragraph_selection(
                                       doc_path, images_folder, image_name_entry.get(), new_card_window, maps_window),
                                   bg="#007BFF", fg="white", relief="flat")
    emplacement_button.pack(pady=20)

def supprimer_carte(doc_path, repere, list_frame):
    """Supprime une carte sélectionnée du document Word."""
    word_app = win32.Dispatch("Word.Application")
    word_app.Visible = False
    document = word_app.Documents.Open(os.path.normpath(doc_path))

    try:
        for paragraph in document.Paragraphs:
            if f"<<{repere}>>" in paragraph.Range.Text:
                paragraph.Range.Delete()
                break

        document.Save()
        afficher_cartes(list_frame, doc_path)  # Mettre à jour la liste des cartes
        messagebox.showinfo("Carte supprimée", f"La carte avec le repère '{repere}' a été supprimée.")
    finally:
        document.Close(SaveChanges=True)
        word_app.Quit()

def afficher_cartes(list_frame, doc_path):
    """Affiche la liste des cartes présentes dans le document avec leur contexte et repère associé."""
    # Nettoyer la liste existante
    for widget in list_frame.winfo_children():
        widget.destroy()

    # Charger le document
    doc = Document(doc_path)
    paragraphs = list(doc.paragraphs)

    for i, paragraph in enumerate(paragraphs):
        if "<<" in paragraph.text and ">>" in paragraph.text:
            # Extraire le repère de la carte
            repere = paragraph.text.strip("<<>>").strip()

            # Trouver le texte le plus proche au-dessus
            texte_proche = "Non spécifié"
            for j in range(i - 1, -1, -1):
                texte_precedent = paragraphs[j].text.strip()
                if texte_precedent:
                    texte_proche = texte_precedent
                    break

            # Créer une section pour la carte
            frame = tk.Frame(list_frame, bg="#f5f5f5", padx=10, pady=5, bd=1, relief="solid")
            frame.pack(fill="x", padx=10, pady=5)

            # Configurer la grille pour s'étendre
            frame.grid_columnconfigure(0, weight=1)
            frame.grid_columnconfigure(1, weight=1)
            frame.grid_columnconfigure(2, weight=0)  # Pour le bouton de suppression

            # Ajouter les labels pour afficher les informations
            tk.Label(frame, text="Nom de la carte :", font=("Helvetica", 10, "bold"), bg="#f5f5f5").grid(row=0, column=0, sticky="w")
            tk.Label(frame, text=repere, bg="#f5f5f5").grid(row=0, column=1, sticky="ew")

            tk.Label(frame, text="Contexte proche :", font=("Helvetica", 10, "bold"), bg="#f5f5f5").grid(row=1, column=0, sticky="w")
            tk.Label(frame, text=texte_proche, bg="#f5f5f5").grid(row=1, column=1, sticky="ew")

            # Boutons pour supprimer la carte
            delete_button = tk.Button(frame, text="Supprimer", command=lambda repere=repere: supprimer_carte(doc_path, repere, list_frame), bg="red", fg="white", relief="raised")
            delete_button.grid(row=0, column=2, rowspan=2, padx=5, sticky="e")

