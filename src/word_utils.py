import os
import re
import time
import unicodedata
import win32com.client as win32
from docx import Document
from docx.shared import Inches
from tkinter import messagebox
from config import BASE_IMAGE_PATH
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import normalize_bookmark_name

def paragraph_contains_image(paragraph):
    blip_elements = paragraph._element.xpath('.//a:blip')
    inline_elements = paragraph._element.xpath('.//wp:inline')
    return bool(blip_elements or inline_elements)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def insert_maps_to_doc(doc_path, images_folder, progress_var, step_text, secondary_progress_var, root, height_entry, width_entry):
    step_text.set("Insertion des cartes dans le document Word...")
    doc = Document(doc_path)

    total_paragraphs = len(doc.paragraphs)
    for i, paragraph in enumerate(doc.paragraphs):
        if "<<" in paragraph.text and ">>" in paragraph.text:
            repere = paragraph.text.strip("<<>>").strip()
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if paragraph_contains_image(next_para):
                    delete_paragraph(next_para)
            chemin_image = os.path.join(images_folder, f"{repere}.jpg")
            if os.path.exists(chemin_image):
                new_para = doc.add_paragraph()
                new_para.add_run().add_picture(chemin_image, width=Inches(width_entry), height=Inches(height_entry))
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph._element.addnext(new_para._element)
        secondary_progress_var.set((i + 1) * 100 / total_paragraphs)
        time.sleep(0.05)
        root.update_idletasks()  # Met à jour l'interface graphique
    doc.save(doc_path)
    progress_var.set(30)

def insert_images(doc_path, image_data, progress_var, step_text, secondary_progress_var, root):
    step_text.set("Suppression des anciens signets et images...")
    doc_path = os.path.normpath(doc_path)
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False

    # Vérifier si le fichier Word existe
    if not os.path.exists(doc_path):
        messagebox.showerror("Erreur", f"Le fichier Word '{doc_path}' est introuvable.")
        return

    # Ouvrir le document Word
    document = word.Documents.Open(doc_path)
    
    total_bookmarks = len(list(document.Bookmarks))
    total_images = len(image_data)
    total_steps = total_bookmarks + total_images  # Combinaison des signets et images pour la barre de progression

    # Supprimer tous les signets et les images associées
    bookmarks_to_delete = list(document.Bookmarks)

    for i, bookmark in enumerate(bookmarks_to_delete):
        if bookmark.Range.InlineShapes.Count > 0:
            bookmark.Range.InlineShapes[0].Delete()  # Supprimer l'image associée
        if bookmark in document.Bookmarks:
            try:
                bookmark.Delete()  # Supprimer le signet
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de la suppression du signet : {e}")

        # Mise à jour de la barre de progression secondaire
        secondary_progress_var.set((i + 1) * 100 / total_steps)
        root.update_idletasks()

    # Mise à jour du texte d'étape
    step_text.set("Insertion des graphiques/tableaux dans le document Word...")

    # Aller à la page 5
    word.Selection.GoTo(What=1, Which=1, Count=5)  # 1 correspond à wdGoToPage et wdGoToAbsolute
    start_range = word.Selection.Range.Start

    for j, data in enumerate(image_data):
        target_text_or_bookmark = data['target']
        image_path = os.path.normpath(os.path.join(BASE_IMAGE_PATH, data['image_path']))
        bookmark_name = normalize_bookmark_name(data['bookmark_name'])

        # Vérifier si l'image existe
        if not os.path.exists(image_path):
            messagebox.showerror("Erreur", f"Le fichier image '{image_path}' est introuvable.")
            continue

        # Rechercher le texte spécifié ou un signet à partir de la page 5
        found = False
        range_to_insert = None

        for paragraph in document.Paragraphs:
            if paragraph.Range.Start >= start_range and target_text_or_bookmark in paragraph.Range.Text:
                found = True
                range_to_insert = paragraph.Range
                break

        if not found:
            target_text_or_bookmark = normalize_bookmark_name(target_text_or_bookmark)
            for bookmark in document.Bookmarks:
                if bookmark.Name == target_text_or_bookmark and bookmark.Range.Start >= start_range:
                    found = True
                    range_to_insert = bookmark.Range
                    break

        if found:
            range_to_insert.Collapse(Direction=0)
            new_image = range_to_insert.InlineShapes.AddPicture(image_path)
            new_image.LockAspectRatio = False
            new_image.Width = document.PageSetup.PageWidth - document.PageSetup.LeftMargin - document.PageSetup.RightMargin
            new_image.Height = new_image.Height * 0.8
            new_image.Range.ParagraphFormat.Alignment = 1
            document.Bookmarks.Add(bookmark_name, new_image.Range)

        # Mise à jour de la barre de progression secondaire
        secondary_progress_var.set((total_bookmarks + j + 1) * 100 / total_steps)
        root.update_idletasks()

    # Sauvegarder et fermer le document Word
    document.Save()
    document.Close()
    word.Quit()

    # Mise à jour de la barre de progression principale à 100%
    progress_var.set(100)